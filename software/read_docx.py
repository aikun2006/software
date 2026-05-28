from docx import Document
import sys

doc_path = sys.argv[1]
doc = Document(doc_path)

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Also try to extract tables
for table in doc.tables:
    print("\n=== TABLE ===")
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(" | ".join(row_text))