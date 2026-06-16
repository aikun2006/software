from docx import Document

doc = Document(r'd:\2026比赛\26软件杯相关\software\示范景区公开资料包\灵山胜境 景点结构化数据集.docx')

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text}")
        print("-" * 40)

print("\n\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} (rows: {len(table.rows)}, cols: {len(table.columns)}) ---")
    for j, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ')[:80] for cell in row.cells]
        print(f"Row {j}: {' | '.join(cells)}")