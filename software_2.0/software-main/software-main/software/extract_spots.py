from docx import Document
import sys

doc = Document(r'd:\2026比赛\26软件杯相关\software\示范景区公开资料包\灵山胜境 景点结构化数据集.docx')

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text}")

print("\n\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} (rows: {len(table.rows)}, cols: {len(table.columns)}) ---")
    for j, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        # Print header
        if j == 0:
            print(f"Header: {cells}")
        else:
            print(f"\nRow {j}:")
            for k, cell in enumerate(cells):
                if cell:
                    print(f"  Col {k}: {cell[:200]}..." if len(cell) > 200 else f"  Col {k}: {cell}")